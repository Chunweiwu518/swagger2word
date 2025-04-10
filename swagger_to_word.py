# Requirements: python-docx PyYAML

import json
import yaml
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def add_heading(doc, text, level=1):
    """Adds a heading to the document."""
    try:
        doc.add_heading(text, level=level)
    except Exception as e:
        print(f"Error adding heading '{text}' (level {level}): {e}")
        # Fallback to paragraph if heading fails
        add_paragraph(doc, text, bold=True)


def add_paragraph(doc, text, bold=False, italic=False, style=None):
    """Adds a paragraph to the document."""
    try:
        # Replace None or invalid characters with empty string or placeholder
        if text is None:
            text = ""
        # Basic sanitization for control characters (optional, adjust as needed)
        text = "".join(c for c in str(text) if ord(c) >= 32 or c in ('\t', '\n', '\r'))

        p = doc.add_paragraph(style=style)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p
    except Exception as e:
        print(f"Error adding paragraph with text '{str(text)[:50]}...': {e}")
        # Attempt to add a placeholder paragraph
        try:
            doc.add_paragraph(f"[Error adding content: {e}]")
        except:
            pass # Ignore if even adding the error message fails


def add_code_block(doc, text):
    """Adds a paragraph formatted as a code block."""
    try:
        if text is None:
            text = ""
        text = "".join(c for c in str(text) if ord(c) >= 32 or c in ('\t', '\n', '\r'))

        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    except Exception as e:
         print(f"Error adding code block with text '{str(text)[:50]}...': {e}")
         try:
            doc.add_paragraph(f"[Error adding code block: {e}]")
         except:
            pass


def format_schema(schema, indent=0):
    """Recursively formats a schema object into a readable string list."""
    indent_str = "  " * indent
    lines = []
    if not isinstance(schema, dict):
        # Handle non-dict schema like 'anyOf', 'allOf', etc. if needed, or return error/string representation
        if isinstance(schema, list): # Handle cases like anyOf/allOf which might be lists
             lines.append(f"{indent_str}Schema list:")
             for item in schema:
                  lines.extend(format_schema(item, indent + 1))
             return lines
        elif schema is None:
             return [f"{indent_str}Schema is null"]
        else:
             return [f"{indent_str}Invalid schema format: {type(schema)} - Content: {str(schema)[:100]}"] # Show part of the content

    schema_type = schema.get('type', 'any') # Default to 'any' if type is missing

    # Add description first if available at the current level
    description = schema.get('description')
    if description:
        lines.append(f"{indent_str}Description: {description}")

    # Add title if available
    title = schema.get('title')
    if title:
        lines.append(f"{indent_str}Title: {title}")

    # Handle $ref
    if '$ref' in schema:
        lines.append(f"{indent_str}Reference: {schema['$ref']}")
        # Optionally add other properties if they exist alongside $ref (less common)
        for key, value in schema.items():
            if key not in ['$ref', 'description', 'title']:
                lines.append(f"{indent_str}{key}: {value}") # Display other properties simply

    # Handle composition keywords (allOf, anyOf, oneOf)
    elif any(key in schema for key in ['allOf', 'anyOf', 'oneOf']):
        for key in ['allOf', 'anyOf', 'oneOf']:
            if key in schema:
                lines.append(f"{indent_str}{key}:")
                if isinstance(schema[key], list):
                    for sub_schema in schema[key]:
                        lines.extend(format_schema(sub_schema, indent + 1))
                else:
                    lines.append(f"{indent_str}  Invalid '{key}' format: Expected list, got {type(schema[key])}")

    # Handle standard types (object, array, basic types)
    elif schema_type == 'object':
        lines.append(f"{indent_str}Type: object")
        # Required properties
        required_props = schema.get('required', [])
        if required_props:
             lines.append(f"{indent_str}Required: {', '.join(required_props)}")

        # Properties
        if 'properties' in schema and isinstance(schema['properties'], dict):
            lines.append(f"{indent_str}Properties:")
            for prop_name, prop_schema in schema['properties'].items():
                req = ' (required)' if prop_name in required_props else ''
                lines.append(f"{indent_str}  - {prop_name}{req}:")
                lines.extend(format_schema(prop_schema, indent + 2))
        elif 'properties' in schema:
             lines.append(f"{indent_str}Properties: Invalid format ({type(schema['properties'])})")

        # Additional Properties
        additional_props = schema.get('additionalProperties')
        if additional_props is not None: # Can be boolean or schema object
             lines.append(f"{indent_str}Additional Properties:")
             if isinstance(additional_props, dict):
                 lines.extend(format_schema(additional_props, indent + 2))
             else:
                 lines.append(f"{indent_str}  Allowed: {additional_props}") # Typically boolean

    elif schema_type == 'array':
        lines.append(f"{indent_str}Type: array")
        if 'items' in schema:
            lines.append(f"{indent_str}Items:")
            lines.extend(format_schema(schema['items'], indent + 1))
        else:
            lines.append(f"{indent_str}Items: (Not specified)") # Indicate if 'items' is missing

    else: # Basic types (string, number, integer, boolean, null)
        lines.append(f"{indent_str}Type: {schema_type}")
        # Add common constraints/details for basic types
        if 'format' in schema:
            lines.append(f"{indent_str}Format: {schema.get('format')}")
        if 'enum' in schema:
            lines.append(f"{indent_str}Enum: {', '.join(map(str, schema.get('enum', [])))}")
        if 'pattern' in schema:
            lines.append(f"{indent_str}Pattern: {schema.get('pattern')}")
        if 'minLength' in schema:
            lines.append(f"{indent_str}MinLength: {schema.get('minLength')}")
        if 'maxLength' in schema:
             lines.append(f"{indent_str}MaxLength: {schema.get('maxLength')}")
        if 'minimum' in schema:
            lines.append(f"{indent_str}Minimum: {schema.get('minimum')}")
        if 'maximum' in schema:
            lines.append(f"{indent_str}Maximum: {schema.get('maximum')}")
        if 'default' in schema:
            lines.append(f"{indent_str}Default: {schema.get('default')}")

    # Add example if available at the current level and not handled by basic types above
    if 'example' in schema:
        # Avoid duplicating example if already shown in basic types? No, API spec can have separate examples.
        example_val = schema.get('example')
        # Basic formatting for example
        if isinstance(example_val, (dict, list)):
             try:
                 example_str = json.dumps(example_val, indent=2, ensure_ascii=False)
                 # Corrected approach: Split lines and indent subsequent lines properly
                 example_lines = example_str.splitlines()
                 # The first line should align with "Example:", indented further by indent_str
                 # Subsequent lines should also be indented by indent_str
                 indented_example_str = "\n".join([f"{indent_str}{line}" for line in example_lines])
                 lines.append(f"{indent_str}Example:\n{indented_example_str}")
             except Exception as e: # Catch potential JSON errors
                 print(f"Warning: Could not format example value: {e}")
                 lines.append(f"{indent_str}Example: {example_val} (raw)")
        else:
             lines.append(f"{indent_str}Example: {example_val}")

    return lines


def parse_and_write_doc(filepath):
    """Parses a single Swagger/OpenAPI file and writes its content to a separate docx file."""
    print(f"Processing {filepath}...")
    base_name = os.path.splitext(os.path.basename(filepath))[0]
    output_filename = f"{base_name}_api_doc.docx"

    document = Document() # Create a new document for this file

    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            if filepath.lower().endswith('.json'):
                spec = json.load(f)
            elif filepath.lower().endswith(('.yaml', '.yml')):
                spec = yaml.safe_load(f)
            else:
                print(f"Skipping unsupported file format: {filepath}")
                return False # Indicate failure
    except Exception as e:
        print(f"Error reading or parsing {filepath}: {e}")
        # Create a simple error doc for this file
        add_heading(document, f"Error processing: {os.path.basename(filepath)}", level=1)
        add_paragraph(document, f"Could not read or parse the file. Error: {e}")
        try:
            document.save(output_filename)
            print(f"Saved error report to '{output_filename}'")
        except Exception as save_e:
            print(f"Could not save error report '{output_filename}': {save_e}")
        return False # Indicate failure

    try:
        # --- Add Title ---
        document.add_heading(f"API Specification: {os.path.basename(filepath)}", level=0) # Use level 0 for main title

        # --- General Info ---
        if 'info' in spec:
            info = spec['info']
            add_heading(document, "General Information", level=1) # Level 1 for sections
            add_paragraph(document, f"Title: {info.get('title', 'N/A')}", bold=True)
            add_paragraph(document, f"Version: {info.get('version', 'N/A')}")
            if 'description' in info:
                add_paragraph(document, f"Description: {info.get('description')}")

        # --- Servers ---
        if 'servers' in spec and spec['servers']:
            add_heading(document, "Servers", level=1)
            for server in spec['servers']:
                url = server.get('url', 'N/A')
                desc = server.get('description', '')
                add_paragraph(document, f"- URL: {url}{' (' + desc + ')' if desc else ''}")

        # --- Paths (API Endpoints) ---
        if 'paths' in spec:
            add_heading(document, "API Endpoints", level=1)
            for path, path_item in spec.get('paths', {}).items():
                if not isinstance(path_item, dict): continue

                path_params_list = path_item.get('parameters', [])

                for method, operation in path_item.items():
                    if method.upper() not in ['GET', 'POST', 'PUT', 'PATCH', 'DELETE', 'OPTIONS', 'HEAD', 'TRACE'] or not isinstance(operation, dict):
                        continue

                    operation_id = operation.get('operationId', 'N/A')
                    summary = operation.get('summary', '')
                    description = operation.get('description', '')
                    tags = operation.get('tags', [])

                    add_heading(document, f"{method.upper()} {path}", level=2) # Level 2 for each endpoint
                    add_paragraph(document, f"Operation ID: {operation_id}", italic=True)
                    if tags:
                         add_paragraph(document, f"Tags: {', '.join(tags)}", italic=True)
                    if summary:
                        add_paragraph(document, summary, bold=True)
                    if description:
                        add_paragraph(document, description)

                    # --- Parameters ---
                    op_params_list = operation.get('parameters', [])
                    all_params_list = path_params_list + op_params_list
                    unique_params_dict = {}
                    for param in all_params_list:
                       if isinstance(param, dict) and 'name' in param and 'in' in param:
                           key = (param.get('name'), param.get('in'))
                           unique_params_dict[key] = param
                    unique_params = list(unique_params_dict.values())

                    if unique_params:
                        add_heading(document, "Parameters", level=3) # Level 3 for sub-sections
                        try:
                            table = document.add_table(rows=1, cols=5)
                            table.style = 'Table Grid'
                            table.autofit = False
                            # Set column widths (example, adjust as needed)
                            table.columns[0].width = Inches(1.2)
                            table.columns[1].width = Inches(0.6)
                            table.columns[2].width = Inches(0.8)
                            table.columns[3].width = Inches(2.5)
                            table.columns[4].width = Inches(3.0)


                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Name'
                            hdr_cells[1].text = 'In'
                            hdr_cells[2].text = 'Required'
                            hdr_cells[3].text = 'Description'
                            hdr_cells[4].text = 'Schema / Type'

                            for cell in hdr_cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                            for param in unique_params:
                                row_cells = table.add_row().cells
                                row_cells[0].text = param.get('name', 'N/A')
                                row_cells[1].text = param.get('in', 'N/A')
                                row_cells[2].text = str(param.get('required', False))
                                row_cells[3].text = param.get('description', '')
                                schema = param.get('schema', {})
                                schema_text = '\n'.join(format_schema(schema)) if schema else param.get('type', 'N/A') # OpenAPI v2 'type' fallback
                                row_cells[4].text = schema_text
                        except Exception as e:
                            print(f"  Error creating parameters table for {method.upper()} {path}: {e}")
                            add_paragraph(document, "[Error displaying parameters table]", italic=True)

                    # --- Request Body ---
                    if 'requestBody' in operation:
                        req_body = operation['requestBody']
                        if isinstance(req_body, dict):
                            add_heading(document, "Request Body", level=3)
                            add_paragraph(document, f"Required: {req_body.get('required', False)}")
                            if 'description' in req_body:
                                add_paragraph(document, f"Description: {req_body.get('description')}")
                            if 'content' in req_body and isinstance(req_body['content'], dict):
                                for media_type, media_spec in req_body['content'].items():
                                    if isinstance(media_spec, dict):
                                        add_paragraph(document, f"Content-Type: {media_type}", bold=True)
                                        if 'schema' in media_spec:
                                            schema_lines = format_schema(media_spec['schema'])
                                            add_code_block(document, '\n'.join(schema_lines))
                                        if 'examples' in media_spec and isinstance(media_spec['examples'], dict): # Handle multiple examples
                                            add_paragraph(document, "Examples:", bold=True)
                                            for ex_name, ex_value in media_spec['examples'].items():
                                                 add_paragraph(document, f"- {ex_name}:")
                                                 ex_data = ex_value.get('value', ex_value) # Get 'value' if present
                                                 try:
                                                     example_str = json.dumps(ex_data, indent=2, ensure_ascii=False)
                                                     add_code_block(document, example_str)
                                                 except Exception:
                                                     add_code_block(document, str(ex_data))

                                        elif 'example' in media_spec: # Handle single example
                                            add_paragraph(document, "Example:", bold=True)
                                            try:
                                                example_str = json.dumps(media_spec['example'], indent=2, ensure_ascii=False)
                                                add_code_block(document, example_str)
                                            except Exception:
                                                add_code_block(document, str(media_spec['example']))

                    # --- Responses ---
                    if 'responses' in operation and isinstance(operation['responses'], dict):
                        add_heading(document, "Responses", level=3)
                        try:
                            table = document.add_table(rows=1, cols=3)
                            table.style = 'Table Grid'
                            table.autofit = False
                            # Set column widths (example, adjust as needed)
                            table.columns[0].width = Inches(1.0)
                            table.columns[1].width = Inches(2.5)
                            table.columns[2].width = Inches(4.0)

                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Status Code'
                            hdr_cells[1].text = 'Description'
                            hdr_cells[2].text = 'Content / Schema'

                            for cell in hdr_cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                            for status_code, response in operation['responses'].items():
                                if not isinstance(response, dict): continue

                                row_cells = table.add_row().cells
                                row_cells[0].text = str(status_code)
                                row_cells[1].text = response.get('description', '')

                                content_text = []
                                if 'content' in response and isinstance(response['content'], dict):
                                    for media_type, media_spec in response['content'].items():
                                        if isinstance(media_spec, dict):
                                           content_text.append(f"Content-Type: {media_type}")
                                           if 'schema' in media_spec:
                                               schema_lines = format_schema(media_spec['schema'])
                                               content_text.extend([f"  {line}" for line in schema_lines])
                                           # Add example handling for responses if needed, similar to requestBody

                                elif 'schema' in response: # OpenAPI v2 fallback
                                    schema_lines = format_schema(response['schema'])
                                    content_text.extend(schema_lines)

                                row_cells[2].text = '\n'.join(content_text) if content_text else ''
                        except Exception as e:
                            print(f"  Error creating responses table for {method.upper()} {path}: {e}")
                            add_paragraph(document, "[Error displaying responses table]", italic=True)

        # --- Components / Definitions (Schemas) ---
        schemas = None
        if 'components' in spec and isinstance(spec.get('components'), dict) and 'schemas' in spec['components']:
            schemas = spec['components']['schemas']
            add_heading(document, "Schemas (Components)", level=1)
        elif 'definitions' in spec: # OpenAPI v2 fallback
            schemas = spec['definitions']
            add_heading(document, "Definitions", level=1)

        if schemas and isinstance(schemas, dict):
            for schema_name, schema_def in schemas.items():
                add_heading(document, schema_name, level=2) # Level 2 for each schema
                schema_lines = format_schema(schema_def)
                add_code_block(document, '\n'.join(schema_lines))

        # --- Save the document for this file ---
        try:
            document.save(output_filename)
            print(f"Successfully generated '{output_filename}'")
            return True # Indicate success
        except Exception as e:
            print(f"Error saving Word document '{output_filename}': {e}")
            return False # Indicate failure

    except Exception as e:
        print(f"!!! Critical error processing content of {filepath}: {e}")
        # Try to add an error message to the doc even if parsing failed mid-way
        try:
             add_heading(document, f"Critical Error in: {os.path.basename(filepath)}", level=1)
             add_paragraph(document, f"An unexpected error occurred while generating documentation for this file: {e}")
             document.save(output_filename) # Attempt to save the partial doc with error
             print(f"Saved partial document with error to '{output_filename}'")
        except Exception as save_e:
            print(f"Could not save error report or partial document '{output_filename}': {save_e}")
        return False # Indicate failure


# --- Main Script ---
# List of your API specification files
api_files = [
    "assistant-swagger-spec.json",
    "file_manager_swagger.json",
    "history-swagger-spec.json",
    "pegasusiam_swagger.yaml" # Make sure this file exists
]

files_processed_count = 0
files_failed_count = 0

# --- Process each file individually ---
for api_file in api_files:
    if os.path.exists(api_file):
        if parse_and_write_doc(api_file):
            files_processed_count += 1
        else:
            files_failed_count += 1
    else:
        print(f"Warning: File not found, skipping: {api_file}")
        files_failed_count += 1

# --- Final Summary ---
print("\n--- Processing Complete ---")
if files_processed_count > 0:
    print(f"Successfully generated {files_processed_count} Word document(s).")
if files_failed_count > 0:
    print(f"{files_failed_count} file(s) encountered errors or were not found.")
